import streamlit as st
import numpy as np
from langchain_core.messages import HumanMessage
from langchain_groq import ChatGroq   
from dotenv import load_dotenv
import os
import io
from docx import Document
import PyPDF2
from PIL import Image
import pytesseract
import pandas as pd
from st_img_pastebutton import paste
import base64
import cv2

load_dotenv() 

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
st.set_page_config(page_title="AI dịch tiếng Nhật", page_icon="🤖", layout = "centered")
st.title("AI dịch")
st.markdown("Ứng dụng dịch sử dụng mô hình ngôn ngữ lớn (LLM) và các công cụ tích hợp.")

ocr_mode = st.selectbox(
    "Chọn chế độ OCR",
    options=["Dịch ngang (jpn)", "Dịch dọc (jpn_vert)"],
    index=0
)

if "dọc" in ocr_mode.lower():
    ocr_lang = "jpn_vert"
    custom_oem_psm_config = r'--oem 3 --psm 5'
else:
    ocr_lang = "jpn"
    custom_oem_psm_config = r'--oem 3 --psm 6'


model = ChatGroq(
        groq_api_key=os.getenv("GROQ_API_KEY"),
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        temperature=0
    )

uploaded_file = st.file_uploader(
    "Tải lên file văn bản/ảnh tiếng Nhật (txt, docx, pdf, png, jpg, jpeg)", 
    type=["txt", "docx", "pdf", "png", "jpg", "jpeg"]
)
def sharpen_image_pil(pil_img):
    img = np.array(pil_img)
    img = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)

    gaussian = cv2.GaussianBlur(img, (9, 9), 10)

    sharpened = cv2.addWeighted(img, 1.5, gaussian, -0.5, 0)

    sharpened_rgb = cv2.cvtColor(sharpened, cv2.COLOR_BGR2RGB)
    return Image.fromarray(sharpened_rgb)

pasted_file = paste(label="📋 Click here, then paste nội dung/ảnh", key="pastebox")


def read_file(file_input):
    text = ""

    if hasattr(file_input, "type"):
        if file_input.type == "text/plain":
            text = file_input.read().decode("utf-8")
        elif file_input.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file_input)
            text = "\n".join([p.text for p in doc.paragraphs])
        elif file_input.type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(file_input)
            text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif file_input.type in ["image/png", "image/jpeg"]:
            image = Image.open(file_input)
            image = sharpen_image_pil(image)
            st.image(image, caption="📷 Ảnh đã tải lên", use_container_width=True)
            text = pytesseract.image_to_string(image, lang=ocr_lang, config=custom_oem_psm_config)

    elif isinstance(file_input, str):
        header, encoded = file_input.split(",", 1)
        binary_data = base64.b64decode(encoded)
        image = Image.open(io.BytesIO(binary_data)).convert("RGB")
        image = sharpen_image_pil(image)
        st.image(image, caption="📷 Ảnh đã dán", use_container_width=True)
        text = pytesseract.image_to_string(image, lang=ocr_lang, config=custom_oem_psm_config)

    elif isinstance(file_input, Image.Image):
        st.image(file_input, caption="📷 Ảnh đã dán", use_container_width=True)
        text = pytesseract.image_to_string(file_input, lang=ocr_lang)

    else:
        st.warning("❌ Không nhận dạng được file/ảnh.")

    return text



if uploaded_file is not None or pasted_file is not None:
    if uploaded_file is not None:
        japanese_text = read_file(uploaded_file)
    else:
        japanese_text = read_file(pasted_file)

    st.subheader("📄 Văn bản gốc (tiếng Nhật OCR/Text)")
    st.text_area("Nội dung", japanese_text, height=200)

    if st.button("Dịch sang Tiếng Việt"):
        response = model.invoke([
    HumanMessage(
        content=(
            "You are an experienced Japanese language teacher. "
            "The student has a JLPT N4 level. Please translate the following Japanese text into Vietnamese "
            "**accurately and fluently**.\n\n"
            "**Output format:**\n"
            "1. First, provide the full Vietnamese translation of the paragraph.\n"
            "2. Then present a table with three columns (Kanji | Hiragana | Vietnamese meaning) for the main vocabulary in the text.\n"
            "3. After each translated sentence, include brief notes on any relevant grammar points or usage as needed.\n\n"
            f"Text: {japanese_text}"
        ),
        additional_kwargs={"job_role": "Japanese language teacher"}
    )
])

        
        st.subheader("📖 Bản dịch & Chú thích")
        result_text = response.content
        st.write(result_text)

        try:
            df = pd.read_csv(io.StringIO(result_text), sep="|").dropna(axis=1, how="all")
            st.dataframe(df, use_container_width=True)
        except Exception:
            st.info("👉 Không phân tích được bảng, hiển thị dạng text.")

        with io.BytesIO() as buffer:
            doc = Document()
            doc.add_heading("Bản dịch & Chú thích", level=1)
            doc.add_paragraph(result_text)
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="⬇️ Tải kết quả (.docx)",
                data=buffer,
                file_name="ket_qua_dich.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
